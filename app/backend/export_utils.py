from io import BytesIO

from docx import Document
from openpyxl import Workbook
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph


# ==========================================================
# Word Export
# ==========================================================

def export_to_docx(title: str, content: str):

    document = Document()

    document.add_heading(title, level=1)

    document.add_paragraph(content)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer


# ==========================================================
# PDF Export
# ==========================================================

def export_to_pdf(title: str, content: str):

    buffer = BytesIO()

    document = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = [
        Paragraph(f"<b>{title}</b>", styles["Heading1"]),
        Paragraph(content.replace("\n", "<br/>"), styles["BodyText"])
    ]

    document.build(story)

    buffer.seek(0)

    return buffer

# ==========================================================
# Excel Export
# ==========================================================

def export_to_excel(title: str, content: str):

    workbook = Workbook()

    sheet = workbook.active

    sheet.title = "AI QA Output"

    # Title
    sheet.append([title])
    sheet.append([])

    # Write each line into Excel
    for line in content.split("\n"):
        sheet.append([line])

    buffer = BytesIO()

    workbook.save(buffer)

    buffer.seek(0)

    return buffer
